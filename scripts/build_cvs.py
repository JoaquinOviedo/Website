"""Build editable, ATS-friendly DOCX CVs from cv_content.py."""

from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

from cv_content import CONTACT, CV, PROFILE_IMAGE, ROOT

OUT = ROOT / "documents" / "cv"
TMP = ROOT / "tmp" / "cv-build"
OUT.mkdir(parents=True, exist_ok=True)
TMP.mkdir(parents=True, exist_ok=True)
FONT = "Arial"
INK = RGBColor(24, 24, 24)
MUTED = RGBColor(70, 70, 70)


def set_font(run, size=10.0, bold=False, italic=False, color=INK):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    return run


def add_link(paragraph, label, url, size=9.0):
    rel = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    for tag, value in (("w:rFonts", None), ("w:color", "202020"), ("w:u", "single"), ("w:sz", str(int(size * 2)))):
        node = OxmlElement(tag)
        if tag == "w:rFonts":
            node.set(qn("w:ascii"), FONT); node.set(qn("w:hAnsi"), FONT)
        elif tag == "w:u": node.set(qn("w:val"), value)
        else: node.set(qn("w:val"), value)
        props.append(node)
    run.append(props)
    text = OxmlElement("w:t"); text.text = label; run.append(text)
    hyperlink.append(run); paragraph._p.append(hyperlink)


def bottom_rule(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    border = OxmlElement("w:bottom")
    for key, value in (("w:val", "single"), ("w:sz", "5"), ("w:space", "2"), ("w:color", "777777")):
        border.set(qn(key), value)
    borders.append(border); ppr.append(borders)


def crop_photo():
    target = TMP / "profile-square.jpg"
    image = Image.open(PROFILE_IMAGE).convert("RGB")
    side = min(image.size)
    left = (image.width - side) // 2
    top = max(0, min((image.height - side) // 2 - 70, image.height - side))
    image.crop((left, top, left + side, top + side)).resize((900, 900), Image.Resampling.LANCZOS).save(target, quality=92, optimize=True)
    return target


def configure(doc):
    section = doc.sections[0]
    section.page_width = Mm(210); section.page_height = Mm(297)
    section.top_margin = Inches(0.52); section.bottom_margin = Inches(0.48)
    section.left_margin = Inches(0.65); section.right_margin = Inches(0.65)
    normal = doc.styles["Normal"]
    normal.font.name = FONT; normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT); normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.paragraph_format.space_after = Pt(2); normal.paragraph_format.line_spacing = 1.18
    bullet = doc.styles["List Bullet"]
    bullet.font.name = FONT; bullet.font.size = Pt(9.7)
    bullet.paragraph_format.left_indent = Inches(0.20); bullet.paragraph_format.first_line_indent = Inches(-0.14)
    bullet.paragraph_format.space_after = Pt(1.35); bullet.paragraph_format.line_spacing = 1.15
    props = doc.core_properties
    props.title = "Joaquín Nicolás Oviedo - Curriculum Vitae"; props.subject = "Professional resume"
    props.author = ""; props.last_modified_by = ""; props.comments = ""; props.keywords = ""


def section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4.5); p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(text), 10.3, bold=True)
    bottom_rule(p)


def item_heading(doc, title, period, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(0.5)
    set_font(p.add_run(title), 9.9, bold=True)
    set_font(p.add_run(f" | {period}"), 9.4, bold=True, color=MUTED)
    if subtitle:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0.5)
        set_font(p.add_run(subtitle), 9.4, italic=True, color=MUTED)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_font(p.add_run(text), 9.7)


def add_header(doc, data):
    if data["photo"]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(str(crop_photo()), width=Mm(27), height=Mm(27))
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run(CONTACT["name"]), 17.5, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(1.5)
    set_font(p.add_run(data["title"]), 10.1, bold=True, color=MUTED)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(0.8)
    set_font(p.add_run(f'{CONTACT["location"]} | '), 8.9)
    add_link(p, CONTACT["email"], f'mailto:{CONTACT["email"]}')
    set_font(p.add_run(" | "), 8.9); add_link(p, CONTACT["linkedin_label"], CONTACT["linkedin_url"])
    set_font(p.add_run(" | "), 8.9); add_link(p, CONTACT["github_label"], CONTACT["github_url"])
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(1)
    add_link(p, CONTACT["portfolio_label"], data["portfolio_url"], 8.9)
    bottom_rule(p)


def build(locale):
    data = CV[locale]
    doc = Document(); configure(doc); add_header(doc, data)
    section_title(doc, data["sections"]["summary"])
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
    set_font(p.add_run(data["summary"]), 9.8)
    section_title(doc, data["sections"]["experience"])
    exp = data["experience"]
    item_heading(doc, exp["company"], exp["period"], exp["role"])
    for text in exp["bullets"]: add_bullet(doc, text)
    section_title(doc, data["sections"]["projects"])
    for project in data["projects"]:
        item_heading(doc, project["heading"], project["period"])
        for text in project["bullets"]: add_bullet(doc, text)
    section_title(doc, data["sections"]["education"])
    for title, period in data["education"]: item_heading(doc, title, period)
    section_title(doc, data["sections"]["skills"])
    for label, value in data["skills"]:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
        set_font(p.add_run(f"{label}: "), 9.5, bold=True); set_font(p.add_run(value), 9.5)
    section_title(doc, data["sections"]["languages"])
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
    set_font(p.add_run(data["languages"]), 9.4)
    p = doc.add_paragraph(); set_font(p.add_run(data["training"]), 9.4)
    target = OUT / data["filename"]
    doc.save(target)
    return target


if __name__ == "__main__":
    for locale in ("es", "en"):
        print(build(locale))
