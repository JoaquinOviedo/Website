"""Structural and text checks for the generated CV deliverables."""

import re
import sys
import zipfile
from pathlib import Path

import pdfplumber
from docx import Document
from pypdf import PdfReader

from cv_content import CONTACT, CV, ROOT

MAX_BYTES = 2_500_000
REQUIRED = [CONTACT["name"], CONTACT["email"], CONTACT["portfolio_label"], "Circo Studio", "WIRIN"]
FORBIDDEN = ["(cid:", "English: B2", "Inglés B2"]


def fail(message):
    print(f"ERROR: {message}", file=sys.stderr)
    raise SystemExit(1)


def verify_pdf(locale, path):
    if not path.exists() or path.stat().st_size >= MAX_BYTES:
        fail(f"Missing or oversized PDF: {path}")
    reader = PdfReader(path)
    if len(reader.pages) != 1:
        fail(f"PDF must have exactly one page: {path}")
    with pdfplumber.open(path) as pdf:
        text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    for expected in REQUIRED:
        if expected not in text:
            fail(f"Missing '{expected}' in {path}")
    for forbidden in FORBIDDEN:
        if forbidden in text:
            fail(f"Forbidden text '{forbidden}' in {path}")
    image_count = sum(len(page.images) for page in reader.pages)
    if image_count != (1 if locale == "es" else 0):
        fail(f"Unexpected image count in {path}: {image_count}")
    links = sum(len(page.get("/Annots", [])) for page in reader.pages)
    if links < 4:
        fail(f"Expected email and three public links in {path}")
    print(f"OK PDF {locale}: 1 page, {path.stat().st_size} bytes, {links} links")


def verify_docx(locale, path):
    if not path.exists() or path.stat().st_size >= MAX_BYTES:
        fail(f"Missing or oversized DOCX: {path}")
    doc = Document(path)
    if doc.tables:
        fail(f"ATS DOCX must not contain tables: {path}")
    text = "\n".join(p.text for p in doc.paragraphs)
    for expected in REQUIRED:
        if expected not in text:
            fail(f"Missing '{expected}' in {path}")
    for forbidden in FORBIDDEN:
        if forbidden in text:
            fail(f"Forbidden text '{forbidden}' in {path}")
    with zipfile.ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        rels = archive.read("word/_rels/document.xml.rels").decode("utf-8")
    drawings = len(re.findall(r"<w:drawing", document_xml))
    if drawings != (1 if locale == "es" else 0):
        fail(f"Unexpected drawing count in {path}: {drawings}")
    if rels.count("TargetMode=\"External\"") < 4:
        fail(f"Expected email and three public links in {path}")
    print(f"OK DOCX {locale}: no tables, {drawings} image(s), {path.stat().st_size} bytes")


if __name__ == "__main__":
    for locale, data in CV.items():
        verify_pdf(locale, ROOT / "public" / "cv" / data["pdf_filename"])
        verify_docx(locale, ROOT / "documents" / "cv" / data["filename"])
    for data in CV.values():
        public = ROOT / "public" / "cv" / data["pdf_filename"]
        copy = ROOT / "output" / "pdf" / data["pdf_filename"]
        if public.read_bytes() != copy.read_bytes():
            fail(f"Verified copy differs from public PDF: {copy}")
    print("All CV checks passed.")
